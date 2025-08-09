"""
ADGM Corporate Agent Test Document Generator
This script generates various test DOCX files to test the ADGM Corporate Agent system
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_test_directory():
    """Create a directory for test documents"""
    if not os.path.exists('test_documents'):
        os.makedirs('test_documents')
    return 'test_documents'

def generate_articles_of_association_with_issues():
    """Generate Articles of Association with multiple issues for testing"""
    doc = Document()
    
    # Title
    title = doc.add_heading('ARTICLES OF ASSOCIATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('OF', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('XYZ TECHNOLOGY LIMITED', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    date_para = doc.add_paragraph(f'DATED: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Part 1: Preliminary
    doc.add_heading('PART 1: PRELIMINARY', 1)
    
    # Article 1 - Name
    doc.add_heading('Article 1 - Name', 2)
    doc.add_paragraph('The name of the Company is XYZ Technology Limited (the "Company").')
    
    # Article 2 - Registered Office (WITH ISSUE)
    doc.add_heading('Article 2 - Registered Office', 2)
    para = doc.add_paragraph(
        'The registered office of the Company shall be situated at Office 456, '
        'Building B, Dubai International Financial Centre, Dubai, UAE.'
    )
    # ISSUE: References DIFC instead of ADGM
    
    # Article 3 - Jurisdiction (WITH ISSUE)
    doc.add_heading('Article 3 - Jurisdiction', 2)
    doc.add_paragraph(
        'Any disputes arising from these Articles shall be resolved in accordance with '
        'the laws of the United Arab Emirates and the courts of Dubai.'
    )
    # ISSUE: Should reference ADGM Courts and ADGM laws
    
    # Article 4 - Objects
    doc.add_heading('Article 4 - Objects', 2)
    doc.add_paragraph('The objects of the Company are to carry on business in technology and software development.')
    
    # Article 5 - Share Capital
    doc.add_heading('Article 5 - Share Capital', 2)
    doc.add_paragraph(
        'The share capital of the Company shall be USD 100,000 divided into 100,000 '
        'ordinary shares of USD 1 each.'
    )
    
    # Article 6 - Directors (WITH ISSUE)
    doc.add_heading('Article 6 - Directors', 2)
    doc.add_paragraph('The Company shall have minimum one director.')
    # ISSUE: Grammar error - missing "a" or should be "a minimum of"
    
    # Article 7 - Shareholders Meeting (WITH ISSUE)
    doc.add_heading('Article 7 - Shareholders Meeting', 2)
    doc.add_paragraph('Meetings may be called as needed.')
    # ISSUE: Too vague - missing notice periods and procedures
    
    # Article 8 - Transfer of Shares (WITH ISSUE)
    doc.add_heading('Article 8 - Transfer of Shares', 2)
    doc.add_paragraph('Shares are transferable.')
    # ISSUE: Incomplete - missing transfer procedures
    
    # Article 9 - Dividends
    doc.add_heading('Article 9 - Dividends', 2)
    doc.add_paragraph(
        'The Company may declare and pay dividends in accordance with the respective '
        'rights of the shareholders.'
    )
    
    # Article 10 - Winding Up (WITH ISSUE)
    doc.add_heading('Article 10 - Winding Up', 2)
    # ISSUE: Missing content for this section
    
    # Part 2: Signatory Section (WITH ISSUE)
    doc.add_page_break()
    doc.add_heading('PART 2: SIGNATORY SECTION', 1)
    doc.add_paragraph('Signed by the shareholders:')
    doc.add_paragraph()
    doc.add_paragraph('_______________________')
    doc.add_paragraph('Name:')
    doc.add_paragraph('Date:')
    # ISSUE: Incomplete signatory section
    
    return doc

def generate_board_resolution_correct():
    """Generate a correct Board Resolution document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('BOARD RESOLUTION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('OF', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('ALPHA INNOVATIONS LIMITED', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('(A company incorporated in Abu Dhabi Global Market)')
    
    date_para = doc.add_paragraph(f'DATED: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content
    doc.add_paragraph()
    doc.add_paragraph(
        'At a meeting of the Board of Directors of Alpha Innovations Limited '
        '(the "Company") duly convened and held at the registered office of the Company '
        'at Abu Dhabi Global Market Square, Al Maryah Island, Abu Dhabi, UAE, '
        f'on {datetime.now().strftime("%B %d, %Y")} at 10:00 AM.'
    )
    
    doc.add_paragraph()
    doc.add_heading('PRESENT:', 2)
    doc.add_paragraph('1. John Smith - Chairman')
    doc.add_paragraph('2. Sarah Johnson - Director')
    doc.add_paragraph('3. Michael Brown - Director')
    
    doc.add_paragraph()
    doc.add_heading('IN ATTENDANCE:', 2)
    doc.add_paragraph('Emma Wilson - Company Secretary')
    
    doc.add_paragraph()
    doc.add_heading('RESOLUTIONS:', 2)
    
    doc.add_paragraph('IT WAS RESOLVED THAT:')
    
    doc.add_paragraph(
        '1. The Company shall open a corporate bank account with First Abu Dhabi Bank '
        'for the purpose of conducting the Company\'s business operations.'
    )
    
    doc.add_paragraph(
        '2. The authorized signatories for the bank account shall be any two of the '
        'following directors signing jointly:'
    )
    doc.add_paragraph('   a) John Smith')
    doc.add_paragraph('   b) Sarah Johnson')
    doc.add_paragraph('   c) Michael Brown')
    
    doc.add_paragraph(
        '3. The Company Secretary is hereby authorized to complete all necessary '
        'documentation and formalities with the bank.'
    )
    
    doc.add_paragraph(
        '4. This resolution shall be filed with the ADGM Registration Authority '
        'as required under the ADGM Companies Regulations 2020.'
    )
    
    # Signatures
    doc.add_page_break()
    doc.add_heading('SIGNATURES:', 2)
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________')
    doc.add_paragraph('John Smith')
    doc.add_paragraph('Chairman')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________')
    doc.add_paragraph('Sarah Johnson')
    doc.add_paragraph('Director')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________')
    doc.add_paragraph('Michael Brown')
    doc.add_paragraph('Director')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    
    return doc

def generate_shareholder_resolution_with_issues():
    """Generate Shareholder Resolution with some issues"""
    doc = Document()
    
    title = doc.add_heading('SHAREHOLDER RESOLUTION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('BETA ENTERPRISES LLC', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ISSUE: Missing "OF" and proper formatting
    
    doc.add_paragraph()
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%m/%d/%Y")}')
    # ISSUE: Inconsistent date format
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The undersigned, being all the shareholders of Beta Enterprises LLC, '
        'hereby consent to the following resolutions:'
    )
    
    doc.add_heading('RESOLVED:', 2)
    
    doc.add_paragraph(
        '1. That the company increase its share capital from AED 50,000 to AED 150,000.'
    )
    # ISSUE: Using AED instead of USD (ADGM typically uses USD)
    
    doc.add_paragraph(
        '2. That new shares be issued to the existing shareholders in proportion '
        'to their current holdings.'
    )
    
    doc.add_paragraph(
        '3. That the Articles of Association be amended accordingly.'
    )
    # ISSUE: Vague - doesn't specify which articles or how
    
    doc.add_paragraph(
        '4. Any disputes shall be resolved under UAE Federal Law.'
    )
    # ISSUE: Should reference ADGM law, not UAE Federal Law
    
    # ISSUE: Missing proper signature section
    doc.add_paragraph()
    doc.add_paragraph('Signed:')
    doc.add_paragraph('Shareholders')
    
    return doc

def generate_memorandum_of_association():
    """Generate a Memorandum of Association"""
    doc = Document()
    
    title = doc.add_heading('MEMORANDUM OF ASSOCIATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('OF', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('GAMMA HOLDINGS LIMITED', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('(To be incorporated in Abu Dhabi Global Market)')
    
    doc.add_paragraph()
    doc.add_heading('1. NAME', 2)
    doc.add_paragraph('The name of the Company is GAMMA HOLDINGS LIMITED.')
    
    doc.add_heading('2. REGISTERED OFFICE', 2)
    doc.add_paragraph(
        'The registered office of the Company will be situated in the '
        'Abu Dhabi Global Market, Abu Dhabi, United Arab Emirates.'
    )
    
    doc.add_heading('3. OBJECTS', 2)
    doc.add_paragraph('The objects for which the Company is established are:')
    doc.add_paragraph('(a) To carry on business as a holding company;')
    doc.add_paragraph('(b) To acquire and hold shares in other companies;')
    doc.add_paragraph('(c) To provide management services to subsidiary companies;')
    doc.add_paragraph('(d) To invest in securities and other financial instruments.')
    
    doc.add_heading('4. LIABILITY OF MEMBERS', 2)
    doc.add_paragraph(
        'The liability of the members is limited to the amount, if any, '
        'unpaid on the shares held by them.'
    )
    
    doc.add_heading('5. SHARE CAPITAL', 2)
    doc.add_paragraph(
        'The share capital of the Company is USD 500,000 divided into '
        '500,000 ordinary shares of USD 1 each.'
    )
    
    doc.add_heading('6. INITIAL SHAREHOLDERS', 2)
    doc.add_paragraph('We, the undersigned, wish to form a company pursuant to this Memorandum:')
    
    # Table for shareholders
    doc.add_paragraph()
    doc.add_paragraph('Name: Robert Anderson')
    doc.add_paragraph('Address: Villa 123, Abu Dhabi, UAE')
    doc.add_paragraph('Number of shares: 250,000')
    
    doc.add_paragraph()
    doc.add_paragraph('Name: Lisa Martinez')
    doc.add_paragraph('Address: Apartment 456, Al Reem Island, Abu Dhabi, UAE')
    doc.add_paragraph('Number of shares: 250,000')
    
    doc.add_page_break()
    doc.add_heading('SIGNATURES', 2)
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________')
    doc.add_paragraph('Robert Anderson')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________')
    doc.add_paragraph('Lisa Martinez')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    
    return doc

def generate_incorporation_application():
    """Generate an Incorporation Application Form"""
    doc = Document()
    
    title = doc.add_heading('APPLICATION FOR INCORPORATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('ADGM REGISTRATION AUTHORITY', 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_heading('SECTION 1: COMPANY DETAILS', 1)
    
    doc.add_paragraph('Proposed Company Name: DELTA TECH SOLUTIONS LIMITED')
    doc.add_paragraph('Name Reservation Number: NR-2025-001234')
    doc.add_paragraph('Type of Company: Private Company Limited by Shares')
    doc.add_paragraph('Jurisdiction: Abu Dhabi Global Market')
    
    doc.add_heading('SECTION 2: REGISTERED OFFICE', 1)
    
    doc.add_paragraph('Building: ADGM Square')
    doc.add_paragraph('Floor/Office: Level 15, Office 1501')
    doc.add_paragraph('Street: Al Maryah Island')
    doc.add_paragraph('City: Abu Dhabi')
    doc.add_paragraph('Country: United Arab Emirates')
    
    doc.add_heading('SECTION 3: SHARE CAPITAL', 1)
    
    doc.add_paragraph('Authorized Share Capital: USD 200,000')
    doc.add_paragraph('Issued Share Capital: USD 100,000')
    doc.add_paragraph('Number of Shares: 100,000')
    doc.add_paragraph('Nominal Value per Share: USD 1')
    
    doc.add_heading('SECTION 4: DIRECTORS', 1)
    
    doc.add_paragraph('Director 1:')
    doc.add_paragraph('  Name: James Wilson')
    doc.add_paragraph('  Nationality: British')
    doc.add_paragraph('  Date of Birth: January 15, 1975')
    doc.add_paragraph('  Address: Dubai, UAE')
    
    doc.add_paragraph()
    doc.add_paragraph('Director 2:')
    doc.add_paragraph('  Name: Fatima Al Ahmed')
    doc.add_paragraph('  Nationality: Emirati')
    doc.add_paragraph('  Date of Birth: March 22, 1980')
    doc.add_paragraph('  Address: Abu Dhabi, UAE')
    
    doc.add_heading('SECTION 5: SHAREHOLDERS', 1)
    
    doc.add_paragraph('Shareholder 1:')
    doc.add_paragraph('  Name: James Wilson')
    doc.add_paragraph('  Number of Shares: 60,000')
    doc.add_paragraph('  Percentage: 60%')
    
    doc.add_paragraph()
    doc.add_paragraph('Shareholder 2:')
    doc.add_paragraph('  Name: Fatima Al Ahmed')
    doc.add_paragraph('  Number of Shares: 40,000')
    doc.add_paragraph('  Percentage: 40%')
    
    doc.add_heading('SECTION 6: COMPANY SECRETARY', 1)
    doc.add_paragraph('Name: ABC Corporate Services Limited')
    doc.add_paragraph('License Number: CS-2025-0001')
    doc.add_paragraph('Address: ADGM Square, Abu Dhabi')
    
    doc.add_heading('SECTION 7: DECLARATION', 1)
    doc.add_paragraph(
        'I hereby declare that all information provided in this application is true '
        'and correct to the best of my knowledge and belief.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('_______________________')
    doc.add_paragraph('Authorized Signatory')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    
    return doc

def main():
    """Main function to generate all test documents"""
    
    # Create test directory
    test_dir = create_test_directory()
    
    print("Generating ADGM test documents...")
    print("-" * 50)
    
    # Generate documents
    documents = [
        ("Articles_of_Association_WITH_ISSUES.docx", 
         generate_articles_of_association_with_issues(),
         "Articles with multiple compliance issues"),
        
        ("Board_Resolution_CORRECT.docx", 
         generate_board_resolution_correct(),
         "Properly formatted Board Resolution"),
        
        ("Shareholder_Resolution_WITH_ISSUES.docx", 
         generate_shareholder_resolution_with_issues(),
         "Shareholder Resolution with issues"),
        
        ("Memorandum_of_Association.docx", 
         generate_memorandum_of_association(),
         "Standard Memorandum of Association"),
        
        ("Incorporation_Application.docx", 
         generate_incorporation_application(),
         "Company Incorporation Application")
    ]
    
    # Save all documents
    for filename, doc, description in documents:
        filepath = os.path.join(test_dir, filename)
        doc.save(filepath)
        print(f"✓ Generated: {filename}")
        print(f"  Description: {description}")
        print()
    
    print("-" * 50)
    print(f"All test documents have been generated in the '{test_dir}' folder.")
    print("\nTest Scenarios:")
    print("1. Articles_of_Association_WITH_ISSUES.docx:")
    print("   - Wrong jurisdiction (Dubai instead of ADGM)")
    print("   - Grammar errors")
    print("   - Incomplete clauses")
    print("   - Missing signatory details")
    
    print("\n2. Board_Resolution_CORRECT.docx:")
    print("   - Properly formatted")
    print("   - Correct ADGM references")
    print("   - Complete signatory section")
    
    print("\n3. Shareholder_Resolution_WITH_ISSUES.docx:")
    print("   - Wrong currency (AED instead of USD)")
    print("   - UAE Federal Law reference instead of ADGM")
    print("   - Incomplete signatures")
    
    print("\n4. Memorandum_of_Association.docx:")
    print("   - Standard format for testing")
    
    print("\n5. Incorporation_Application.docx:")
    print("   - Complete application form")
    
    print("\nYou can now upload these documents to test your ADGM Corporate Agent!")

if __name__ == "__main__":
    # Check if python-docx is installed
    try:
        from docx import Document
        main()
    except ImportError:
        print("Error: python-docx is not installed.")
        print("Please install it using: pip install python-docx")
        print("\nAfter installation, run this script to generate test documents.")