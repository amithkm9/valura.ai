#!/usr/bin/env python3
"""
Script to create ADGM test documents for the Corporate Agent system
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_shareholder_resolution():
    """Create the ADGM Shareholder Resolution document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('RESOLUTION OF INCORPORATING SHAREHOLDERS', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company name section
    doc.add_paragraph()
    p = doc.add_paragraph('OF')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('[Insert proposed company name]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red to indicate placeholder
    
    doc.add_paragraph()
    p = doc.add_paragraph('DATED')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    p = doc.add_paragraph('[Insert date]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red to indicate placeholder
    
    doc.add_paragraph()
    
    # Main resolution text
    main_text = doc.add_paragraph(
        'We, the undersigned, being the incorporating shareholders, resolve to incorporate '
        'a private company limited by shares in the Abu Dhabi Global Market under the name of '
    )
    main_text.add_run('[insert proposed company name]').font.color.rgb = RGBColor(255, 0, 0)
    main_text.add_run(
        ' (or any other name approved by ADGM Registration Authority), in accordance with '
        'the applicable regulations and sub-ordinate rules of Abu Dhabi Global Market '
        '(the "Company"). The incorporating shareholders duly adopted the resolution set '
        'forth below on '
    )
    main_text.add_run('[insert date]').font.color.rgb = RGBColor(255, 0, 0)
    main_text.add_run(':')
    
    doc.add_paragraph()
    
    # Resolution statement
    p = doc.add_paragraph()
    p.add_run('IT WAS RESOLVED').bold = True
    p.add_run(', to appoint the officers of the company upon incorporation as follows.')
    
    doc.add_paragraph()
    
    # 1. Authorised Signatories
    doc.add_heading('1) Appointment of Authorised Signatory(ies)', level=2)
    
    # Create table for signatories
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Signing Authority (Jointly/Severally)'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()
    
    # 2. Directors
    doc.add_heading('2) Appointment of Director(s)', level=2)
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Type (Individual/Body Corporate)'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()
    
    # 3. Secretary (Optional)
    doc.add_heading('3) Appointment of Secretary(ies) [Optional]', level=2)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Capacity (Jointly/Severally)'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()
    
    # 4. Adoption of Articles
    doc.add_heading('4) Adoption of Articles', level=2)
    
    p = doc.add_paragraph()
    p.add_run('IT WAS RESOLVED').bold = True
    p.add_run(
        ' that the Company adopts the Articles of Association for the purpose of '
        'incorporation of the Company in the Abu Dhabi Global Market.'
    )
    
    doc.add_paragraph()
    
    # 5. Authorised Share Capital
    doc.add_heading('5) Authorised Share Capital [Optional]', level=2)
    
    p = doc.add_paragraph()
    p.add_run('IT WAS RESOLVED').bold = True
    p.add_run(' that the amount of the authorised share capital of the company shall be as follows: ')
    p.add_run('[xxxx]').font.color.rgb = RGBColor(255, 0, 0)
    p.add_run(' USD')
    
    doc.add_paragraph()
    
    # 6. Share Capital
    doc.add_heading('6) Share Capital', level=2)
    
    p = doc.add_paragraph()
    p.add_run('IT WAS RESOLVED').bold = True
    p.add_run(' that the proposed issued share capital of the company shall be as follows:')
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Share Class Name'
    header_cells[1].text = 'Nominal Value'
    header_cells[2].text = 'Number of Issued Shares'
    header_cells[3].text = 'Amount of Issued Shares'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()
    
    # 7. Shareholders
    doc.add_heading('7) Shareholders', level=2)
    
    doc.add_paragraph('The issued share capital shall be structured as follows.')
    
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "Shareholder's Name"
    header_cells[1].text = 'Share Class Name'
    header_cells[2].text = 'Number of Issued Shares'
    header_cells[3].text = 'Amount paid'
    header_cells[4].text = 'Amount unpaid (if any)'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    doc.add_paragraph()
    
    # 8. Appointment of Signatory for Incorporation
    doc.add_heading('8) Appointment of Signatory for Incorporation Purposes [Optional]', level=2)
    
    p = doc.add_paragraph()
    p.add_run('IT WAS FURTHER RESOLVED').bold = True
    p.add_run(', that ')
    p.add_run('[insert authorised persons name(s)]').font.color.rgb = RGBColor(255, 0, 0)
    p.add_run(
        ' is/are, and each acting alone is, hereby authorized to do and perform any and all '
        'such acts, including execution of any and all documents and certificates, as said '
        'person shall deem necessary or advisable, to carry out the purposes of the foregoing '
        'resolutions to complete the incorporation process with the ADGM RA'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature Section
    doc.add_heading('Signature of Incorporating Shareholders', level=2)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "Shareholder's Name"
    header_cells[1].text = "Shareholder's Signature"
    header_cells[2].text = 'Date'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add empty rows for signatures
    for i in range(1, 3):
        for j in range(3):
            table.rows[i].cells[j].text = '_' * 25
    
    # Save the document
    filename = 'adgm-shareholder-resolution-test.docx'
    doc.save(filename)
    print(f"✅ Created: {filename}")
    return filename

def create_articles_of_association():
    """Create a sample Articles of Association document"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('ARTICLES OF ASSOCIATION', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph('OF')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('[COMPANY NAME] LIMITED')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # Article 1: Company Name
    doc.add_heading('Article 1: Company Name', level=2)
    doc.add_paragraph('The name of the Company is [COMPANY NAME] Limited.')
    
    # Article 2: Registered Office
    doc.add_heading('Article 2: Registered Office', level=2)
    doc.add_paragraph(
        'The registered office of the Company shall be situated at [ADDRESS], '
        'Abu Dhabi Global Market, United Arab Emirates.'
    )
    
    # Article 3: Objects
    doc.add_heading('Article 3: Objects and Powers', level=2)
    doc.add_paragraph(
        'The objects for which the Company is established are to carry on business '
        'as a general commercial company and to do all such things as are incidental '
        'or conducive to the attainment of the above objects.'
    )
    
    # Article 4: Share Capital
    doc.add_heading('Article 4: Share Capital', level=2)
    doc.add_paragraph(
        'The share capital of the Company is USD [AMOUNT] divided into [NUMBER] '
        'ordinary shares of USD [VALUE] each.'
    )
    
    # Article 5: Directors
    doc.add_heading('Article 5: Directors', level=2)
    doc.add_paragraph(
        'The Company shall have a minimum of one (1) director. The directors shall '
        'have all powers necessary for managing the business and affairs of the Company.'
    )
    
    # Article 6: Meetings
    doc.add_heading('Article 6: Meetings', level=2)
    doc.add_paragraph(
        'Annual general meetings shall be held once every calendar year. '
        'Notice of meetings shall be given at least 14 days in advance.'
    )
    
    # Article 7: Governing Law
    doc.add_heading('Article 7: Governing Law', level=2)
    doc.add_paragraph(
        'These Articles of Association and the operations of the Company shall be '
        'governed by the laws and regulations of the Abu Dhabi Global Market.'
    )
    
    # Article 8: Dispute Resolution
    doc.add_heading('Article 8: Dispute Resolution', level=2)
    doc.add_paragraph(
        'Any disputes arising out of or in connection with these Articles shall be '
        'resolved by the ADGM Courts.'
    )
    
    filename = 'adgm-articles-of-association-test.docx'
    doc.save(filename)
    print(f"✅ Created: {filename}")
    return filename

def create_board_resolution():
    """Create a sample Board Resolution document"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('BOARD RESOLUTION', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph('[COMPANY NAME] LIMITED')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph()
    
    # Meeting details
    doc.add_paragraph('Date: [INSERT DATE]')
    doc.add_paragraph('Time: [INSERT TIME]')
    doc.add_paragraph('Venue: [INSERT VENUE]')
    
    doc.add_paragraph()
    
    # Directors present
    doc.add_heading('Directors Present:', level=2)
    doc.add_paragraph('1. [Director Name 1]')
    doc.add_paragraph('2. [Director Name 2]')
    
    doc.add_paragraph()
    
    # Quorum
    doc.add_paragraph('QUORUM: A quorum was present.')
    
    doc.add_paragraph()
    
    # Resolutions
    doc.add_heading('RESOLUTIONS:', level=2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('IT WAS RESOLVED THAT:').bold = True
    
    doc.add_paragraph(
        '1. The company proceed with incorporation in the Abu Dhabi Global Market.'
    )
    
    doc.add_paragraph(
        '2. The Articles of Association as presented be and are hereby approved and adopted.'
    )
    
    doc.add_paragraph(
        '3. The authorized signatories be appointed to act on behalf of the company.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('FURTHER RESOLVED THAT:').bold = True
    
    doc.add_paragraph(
        'Any director of the Company be authorized to execute all necessary documents '
        'to give effect to the above resolutions.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    doc.add_heading('Signatures:', level=2)
    
    doc.add_paragraph('_____________________')
    doc.add_paragraph('Director 1')
    doc.add_paragraph('Date: _______________')
    
    doc.add_paragraph()
    
    doc.add_paragraph('_____________________')
    doc.add_paragraph('Director 2')
    doc.add_paragraph('Date: _______________')
    
    filename = 'adgm-board-resolution-test.docx'
    doc.save(filename)
    print(f"✅ Created: {filename}")
    return filename

def main():
    """Create all test documents"""
    print("Creating ADGM test documents...")
    print("=" * 50)
    
    # Create documents
    files = []
    files.append(create_shareholder_resolution())
    files.append(create_articles_of_association())
    files.append(create_board_resolution())
    
    print("=" * 50)
    print("✅ All test documents created successfully!")
    print("\nYou can now upload these documents to test the ADGM Corporate Agent:")
    for f in files:
        print(f"  - {f}")
    
    print("\nThe system should detect that you're doing company incorporation")
    print("and notify you about missing documents:")
    print("  - Incorporation Application Form")
    print("  - Register of Members and Directors")

if __name__ == "__main__":
    main()