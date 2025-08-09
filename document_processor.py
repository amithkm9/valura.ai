#!/usr/bin/env python3
"""
Document Processor Module - Enhanced Version
Handles document loading, parsing, manipulation, and inline commenting
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from typing import List, Dict, Optional, Tuple
import logging
from pathlib import Path
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process and analyze ADGM legal documents with enhanced commenting"""
    
    def __init__(self):
        self.document = None
        self.document_type = None
        self.document_path = None
        self.issues = []
        self.comments_added = []
        
    def load_document(self, file_path: str) -> bool:
        """Load a Word document for processing"""
        try:
            self.document = Document(file_path)
            self.document_path = file_path
            self.document_type = self._identify_document_type()
            logger.info(f"Loaded document: {file_path}")
            logger.info(f"Identified type: {self.document_type}")
            return True
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            return False
    
    def _identify_document_type(self) -> str:
        """Enhanced document type identification"""
        if not self.document:
            return "unknown"
        
        text = self.get_document_text().lower()
        
        # Priority-based document identification
        # Check for incorporation application first (highest priority)
        if any(phrase in text for phrase in [
            "application for incorporation",
            "incorporation application",
            "adgm registration authority",
            "proposed company name",
            "name reservation number"
        ]):
            return "incorporation_application"
        
        # Check for other document types
        patterns = {
            "articles_of_association": {
                "required": ["articles of association"],
                "optional": ["aoa", "article 1", "share capital", "registered office"],
                "threshold": 2
            },
            "board_resolution": {
                "required": ["board resolution"],
                "optional": ["board of directors", "it was resolved", "directors present", "quorum"],
                "threshold": 2
            },
            "shareholder_resolution": {
                "required": ["shareholder resolution", "resolution of incorporating shareholders"],
                "optional": ["shareholders", "resolved", "incorporating"],
                "threshold": 2
            },
            "memorandum": {
                "required": ["memorandum of association"],
                "optional": ["moa", "objects of the company", "liability of members"],
                "threshold": 2
            },
            "employment_contract": {
                "required": ["employment agreement", "employment contract"],
                "optional": ["employee", "salary", "duties", "termination"],
                "threshold": 2
            },
            "ubo_declaration": {
                "required": ["ultimate beneficial owner", "ubo"],
                "optional": ["beneficial ownership", "declaration"],
                "threshold": 1
            },
            "register": {
                "required": ["register of members", "register of directors"],
                "optional": ["shareholding", "directors register"],
                "threshold": 1
            }
        }
        
        for doc_type, config in patterns.items():
            # Check required phrases
            has_required = any(phrase in text for phrase in config["required"])
            if not has_required:
                continue
            
            # Count optional matches
            optional_matches = sum(1 for phrase in config["optional"] if phrase in text)
            
            if optional_matches >= config.get("threshold", 1):
                return doc_type
        
        return "general_document"
    
    def get_document_text(self) -> str:
        """Extract all text from the document"""
        if not self.document:
            return ""
        
        full_text = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return "\n".join(full_text)
    
    def add_comment_to_paragraph(self, paragraph, comment_text: str, author: str = "ADGM Corporate Agent"):
        """Add a comment to a paragraph (creates a highlighted annotation)"""
        # Since python-docx doesn't support true Word comments,
        # we'll add inline annotations with highlighting
        comment_run = paragraph.add_run(f" [COMMENT: {comment_text}]")
        comment_run.font.color.rgb = RGBColor(255, 0, 0)
        comment_run.font.italic = True
        comment_run.font.size = Pt(9)
        
        # Track that we added this comment
        self.comments_added.append({
            "text": paragraph.text[:50] + "...",
            "comment": comment_text
        })
    
    def check_and_comment_jurisdiction(self) -> List[Dict]:
        """Check for jurisdiction issues and add inline comments"""
        issues = []
        
        # Red flag jurisdictions
        incorrect_jurisdictions = {
            "UAE Federal Courts": "Per ADGM Companies Regulations 2020, Art. 6: Replace with 'ADGM Courts'",
            "Dubai Courts": "Per ADGM Companies Regulations 2020, Art. 6: Use 'ADGM Courts' instead",
            "Abu Dhabi Courts": "Per ADGM Companies Regulations 2020, Art. 6: Should be 'ADGM Courts'",
            "DIFC": "Incorrect jurisdiction - must specify 'Abu Dhabi Global Market (ADGM)'",
            "Dubai International Financial Centre": "Wrong jurisdiction - use 'Abu Dhabi Global Market'",
            "mainland UAE": "Specify 'Abu Dhabi Global Market' for ADGM entities",
            "onshore UAE": "ADGM entities must reference 'Abu Dhabi Global Market'"
        }
        
        for i, paragraph in enumerate(self.document.paragraphs):
            para_text_lower = paragraph.text.lower()
            
            for jurisdiction, comment in incorrect_jurisdictions.items():
                if jurisdiction.lower() in para_text_lower:
                    # Highlight the problematic text
                    for run in paragraph.runs:
                        if jurisdiction.lower() in run.text.lower():
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    # Add comment
                    self.add_comment_to_paragraph(paragraph, comment)
                    
                    issues.append({
                        "paragraph": i,
                        "issue": f"Incorrect jurisdiction reference: '{jurisdiction}'",
                        "severity": "high",
                        "suggestion": comment,
                        "regulation": "ADGM Companies Regulations 2020, Art. 6"
                    })
        
        # Check if ADGM is mentioned at all
        text = self.get_document_text().lower()
        adgm_patterns = ["abu dhabi global market", "adgm"]
        has_adgm = any(pattern in text for pattern in adgm_patterns)
        
        if not has_adgm and self.document_type in ["articles_of_association", "board_resolution", "shareholder_resolution"]:
            # Add comment to first paragraph
            if self.document.paragraphs:
                self.add_comment_to_paragraph(
                    self.document.paragraphs[0],
                    "Missing ADGM jurisdiction - Per ADGM Companies Regulations 2020, Art. 6: Must specify 'Abu Dhabi Global Market'"
                )
            
            issues.append({
                "paragraph": 0,
                "issue": "Missing ADGM jurisdiction reference",
                "severity": "high",
                "suggestion": "Add explicit reference to 'Abu Dhabi Global Market (ADGM)' jurisdiction",
                "regulation": "ADGM Companies Regulations 2020, Art. 6"
            })
        
        return issues
    
    def check_and_comment_weak_language(self) -> List[Dict]:
        """Check for weak language and add inline comments"""
        issues = []
        
        weak_terms = {
            "may": ("shall", "Per ADGM legal drafting standards: Use 'shall' for mandatory obligations"),
            "might": ("shall", "Per ADGM legal drafting standards: Replace with 'shall' for binding effect"),
            "could": ("shall", "Per ADGM legal drafting standards: Use 'shall' for mandatory provisions"),
            "possibly": ("shall", "Ambiguous language - use 'shall' for clarity"),
            "perhaps": ("shall", "Uncertain language - replace with 'shall'"),
            "should consider": ("shall", "Weak obligation - use 'shall' for binding requirements")
        }
        
        for i, paragraph in enumerate(self.document.paragraphs):
            para_text = paragraph.text
            
            for weak, (strong, comment) in weak_terms.items():
                pattern = r'\b' + weak + r'\b'
                if re.search(pattern, para_text, re.IGNORECASE):
                    # Highlight weak terms
                    for run in paragraph.runs:
                        if weak.lower() in run.text.lower():
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    
                    # Add comment
                    self.add_comment_to_paragraph(paragraph, comment)
                    
                    issues.append({
                        "paragraph": i,
                        "issue": f"Weak language detected: '{weak}'",
                        "severity": "medium",
                        "suggestion": f"Replace '{weak}' with '{strong}'",
                        "context": para_text[:100],
                        "regulation": "ADGM legal drafting standards"
                    })
        
        return issues
    
    def check_and_comment_required_sections(self) -> List[Dict]:
        """Check for required sections and add comments for missing ones"""
        issues = []
        text = self.get_document_text().lower()
        
        required_sections = {
            "articles_of_association": {
                "company name": "Per ADGM Companies Regulations 2020, Art. 30: Company name required",
                "registered office": "Per ADGM Companies Regulations 2020, Art. 25: Registered office must be specified",
                "share capital": "Per ADGM Companies Regulations 2020, Art. 12: Share capital details required",
                "directors": "Per ADGM Companies Regulations 2020, Art. 15: Director provisions required",
                "shareholders": "Share ownership structure must be defined",
                "meetings": "Meeting procedures must be specified",
                "governing law": "Per ADGM Companies Regulations 2020, Art. 6: Governing law clause required"
            },
            "board_resolution": {
                "date": "Date of resolution required",
                "directors present": "Attendance record required",
                "quorum": "Quorum confirmation required",
                "resolved": "Resolution language required",
                "signatures": "Director signatures required"
            },
            "shareholder_resolution": {
                "shareholders": "Shareholder details required",
                "shareholdings": "Shareholding structure required",
                "resolved": "Resolution language required",
                "signatures": "Shareholder signatures required"
            },
            "incorporation_application": {
                "company details": "Company information section required",
                "registered office": "ADGM registered office address required",
                "share capital": "Share capital structure required",
                "directors": "Director information required",
                "shareholders": "Shareholder details required"
            }
        }
        
        if self.document_type in required_sections:
            missing_sections = []
            
            for section, comment in required_sections[self.document_type].items():
                if section not in text:
                    missing_sections.append((section, comment))
                    
                    issues.append({
                        "paragraph": -1,
                        "issue": f"Missing required section: '{section}'",
                        "severity": "high",
                        "suggestion": f"Add a section covering '{section}'",
                        "regulation": comment
                    })
            
            # Add a summary comment at the beginning if sections are missing
            if missing_sections and self.document.paragraphs:
                summary_comment = f"Missing {len(missing_sections)} required sections: " + \
                                 ", ".join([s[0] for s in missing_sections[:3]])
                if len(missing_sections) > 3:
                    summary_comment += f" and {len(missing_sections) - 3} more"
                
                self.add_comment_to_paragraph(
                    self.document.paragraphs[0],
                    summary_comment
                )
        
        return issues
    
    def check_and_comment_signatory_sections(self) -> List[Dict]:
        """Check signatory sections and add comments"""
        issues = []
        text = self.get_document_text().lower()
        
        # Check for signature blocks
        signature_indicators = ["signature", "signed", "authorized signatory", "_______"]
        has_signature_section = any(indicator in text for indicator in signature_indicators)
        
        if not has_signature_section:
            # Add comment to last paragraph
            if self.document.paragraphs:
                self.add_comment_to_paragraph(
                    self.document.paragraphs[-1],
                    "Per ADGM execution requirements: Add signature blocks with name, title, and date fields"
                )
            
            issues.append({
                "paragraph": len(self.document.paragraphs) - 1,
                "issue": "Missing signature section",
                "severity": "high",
                "suggestion": "Add proper signature blocks with name, title, and date fields",
                "regulation": "ADGM execution requirements"
            })
        
        # Check for incomplete signatures
        for i, paragraph in enumerate(self.document.paragraphs):
            if "name:" in paragraph.text.lower() and not any(
                char.isalpha() for char in paragraph.text.split("name:")[-1]
            ):
                self.add_comment_to_paragraph(
                    paragraph,
                    "Incomplete signature block - signatory name required"
                )
                
                issues.append({
                    "paragraph": i,
                    "issue": "Incomplete signature block",
                    "severity": "medium",
                    "suggestion": "Complete all signature fields",
                    "regulation": "ADGM documentation standards"
                })
        
        return issues
    
    def perform_comprehensive_review(self) -> List[Dict]:
        """Perform all checks and add inline comments"""
        all_issues = []
        
        # Run all checks with commenting
        all_issues.extend(self.check_and_comment_jurisdiction())
        all_issues.extend(self.check_and_comment_weak_language())
        all_issues.extend(self.check_and_comment_required_sections())
        all_issues.extend(self.check_and_comment_signatory_sections())
        
        # Store issues for later reference
        self.issues = all_issues
        
        return all_issues
    
    def save_reviewed_document(self, output_path: str) -> bool:
        """Save the reviewed document with all comments and highlights"""
        if not self.document:
            return False
        
        try:
            # Create a new document for the reviewed version
            reviewed_doc = Document()
            
            # Add review header
            header = reviewed_doc.add_heading('ADGM COMPLIANCE REVIEW REPORT', 0)
            header.alignment = 1  # Center
            
            # Add metadata
            reviewed_doc.add_paragraph(f"Review Date: {datetime.now().strftime('%B %d, %Y')}")
            reviewed_doc.add_paragraph(f"Document Type: {self.document_type.replace('_', ' ').title()}")
            reviewed_doc.add_paragraph(f"Total Issues Found: {len(self.issues)}")
            
            # Add severity breakdown
            high_issues = sum(1 for i in self.issues if i.get("severity") == "high")
            medium_issues = sum(1 for i in self.issues if i.get("severity") == "medium")
            low_issues = sum(1 for i in self.issues if i.get("severity") == "low")
            
            reviewed_doc.add_paragraph(f"High Severity: {high_issues}")
            reviewed_doc.add_paragraph(f"Medium Severity: {medium_issues}")
            reviewed_doc.add_paragraph(f"Low Severity: {low_issues}")
            
            # Add separator
            reviewed_doc.add_paragraph("=" * 70)
            reviewed_doc.add_heading('REVIEWED DOCUMENT WITH INLINE COMMENTS', 1)
            reviewed_doc.add_paragraph("=" * 70)
            
            # Copy the original document with comments
            for paragraph in self.document.paragraphs:
                reviewed_doc.add_paragraph(paragraph.text)
            
            # Add summary of comments at the end
            if self.comments_added:
                reviewed_doc.add_page_break()
                reviewed_doc.add_heading('COMMENT SUMMARY', 1)
                
                for i, comment_info in enumerate(self.comments_added, 1):
                    reviewed_doc.add_paragraph(f"{i}. Location: {comment_info['text']}")
                    comment_para = reviewed_doc.add_paragraph(f"   Comment: {comment_info['comment']}")
                    comment_para.paragraph_format.left_indent = Inches(0.5)
                    reviewed_doc.add_paragraph("")
            
            # Save the document
            reviewed_doc.save(output_path)
            logger.info(f"Saved reviewed document to: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            return False